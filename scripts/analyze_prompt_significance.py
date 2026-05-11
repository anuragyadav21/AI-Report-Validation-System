#!/usr/bin/env python3
"""AI Report Validation System — statistics: descriptives, figures, RM ANOVA, partial η², Holm post-hoc, regression.

Default workbook: **prompt-worksheet_AB.xlsx**.

Writes a **.docx** beside the workbook by default; no stdout on success. Use **--stdout-only** for terminal.

Statistics: (1) mean & SD by prompt, (2) boxplot, RM line plot, mean±95% CI figures (embedded in .docx),
(3) repeated-measures ANOVA, (4) partial eta-squared, (5) paired Holm tests,
(6) OLS regression with cluster-robust SEs (cluster = scenario).
"""
from __future__ import annotations

import argparse
import io
import os
import re
import sys
import tempfile
import warnings
from dataclasses import dataclass, field
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
_mpl = _ROOT / ".mpl_cache"
_mpl.mkdir(exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(_mpl))

import numpy as np
import pandas as pd
import pingouin as pg
import statsmodels.formula.api as smf
from scipy import stats
from statsmodels.stats.multitest import multipletests

from analyze_rm_anova import build_long_frame, partial_eta_sq

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e

ROOT = _ROOT
DEFAULT_XLSX = ROOT / "prompt-worksheet_AB.xlsx"


@dataclass
class FigurePanel:
    """One figure with heading (H2 in Word), explanatory text, and image path."""

    heading: str
    paragraphs: list[str]
    image_path: Path


@dataclass
class ReportBlock:
    """One report section: prose, optional figure panels, optional data table."""

    title: str
    paragraphs: list[str] = field(default_factory=list)
    panels: list[FigurePanel] = field(default_factory=list)
    table_headers: list[str] | None = None
    table_rows: list[list[str]] | None = None


def _strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def _ascii_table(headers: list[str], rows: list[list[str]]) -> str:
    cols = [list(headers)]
    for r in rows:
        cols.append([str(c) for c in r])
    widths = [max(len(row[i]) for row in cols) for i in range(len(headers))]
    sep = " | ".join("-" * w for w in widths)
    out = []
    out.append(" | ".join(str(headers[i]).ljust(widths[i]) for i in range(len(headers))))
    out.append(sep)
    for r in rows:
        out.append(" | ".join(str(r[i]).ljust(widths[i]) for i in range(len(headers))))
    return "\n".join(out)


def _wide_complete(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    comp = df.dropna(subset=["Overall"]).copy()
    if comp.empty:
        raise ValueError("No numeric Overall scores.")
    wide = comp.pivot_table(
        index="ScenarioID", columns="Prompt", values="Overall", aggfunc="first"
    )
    if not {"A", "B", "C"}.issubset(wide.columns):
        raise ValueError("Need all three prompt levels A, B, C in the workbook.")
    wc = wide.dropna(subset=["A", "B", "C"], how="any")
    if len(wc) < 2:
        raise ValueError("Need at least 2 complete scenarios (non-missing A, B, C).")
    long = wc.reset_index().melt(
        id_vars="ScenarioID",
        var_name="Prompt",
        value_name="Overall",
    )
    return wc, long


def _descriptive_table(long: pd.DataFrame) -> tuple[list[str], list[list[str]]]:
    g = long.groupby("Prompt")["Overall"].agg(["mean", "std"]).round(4)
    headers = ["Prompt", "Mean overall", "SD"]
    rows: list[list[str]] = []
    for prompt in ["A", "B", "C"]:
        if prompt not in g.index:
            continue
        m = float(g.loc[prompt, "mean"])
        s = float(g.loc[prompt, "std"])
        s_str = "—" if (s is None or (isinstance(s, float) and np.isnan(s))) else f"{s:.4f}"
        rows.append([f"Prompt {prompt}", f"{m:.4f}", s_str])
    return headers, rows


def _p_omnibus_from_row(pr: pd.Series) -> float:
    for key in pr.index:
        kl = str(key).lower().replace("-", "")
        if kl in ("punc",) or ("p" in kl and "unc" in kl and "corr" not in kl):
            return float(pr[key])
    raise KeyError("Could not find omnibus p-value in ANOVA row")


def _anova_effect_parts(long: pd.DataFrame) -> tuple[str, str]:
    aov = pg.rm_anova(
        data=long,
        dv="Overall",
        within="Prompt",
        subject="ScenarioID",
        detailed=True,
        correction="auto",
    )
    prompt_row = aov.loc[aov["Source"] == "Prompt"].iloc[0]
    err_row = aov.loc[aov["Source"] == "Error"].iloc[0]

    df1 = int(round(float(prompt_row["DF"])))
    df2 = int(round(float(err_row["DF"])))
    f_val = float(prompt_row["F"])
    p_val = _p_omnibus_from_row(prompt_row)
    p_disp = f"p = {p_val:.4g}" if p_val >= 0.0001 else "p < 0.0001"

    ss_p = float(prompt_row["SS"])
    ss_e = float(err_row["SS"])
    np2 = partial_eta_sq(ss_p, ss_e)

    anova_text = (
        "One-way repeated-measures ANOVA (within factor = Prompt A/B/C, subject = Scenario).\n"
        "Tests: Does prompt type significantly affect Overall scores?\n\n"
        f"F({df1}, {df2}) = {f_val:.4f}, {p_disp}"
    )

    if np2 is not None:
        if np2 >= 0.14:
            interp = "large (conventional η_p² ≥ 0.14)."
        elif np2 >= 0.06:
            interp = "medium (0.06 ≤ η_p² < 0.14)."
        elif np2 >= 0.01:
            interp = "small (0.01 ≤ η_p² < 0.06)."
        else:
            interp = "below the usual small-effect benchmark (η_p² < 0.01)."
        eta_text = (
            f"Partial eta-squared (prompt main effect): η_p² = {np2:.4f}\n"
            f"Magnitude: {interp}"
        )
    else:
        eta_text = "Partial eta-squared could not be computed from the ANOVA sums of squares."

    return anova_text, eta_text


def _posthoc_holm_parts(
    wide: pd.DataFrame,
) -> tuple[list[str], list[str], list[list[str]]]:
    pairs = [
        ("B", "A", "B − A (same scenario)"),
        ("C", "A", "C − A"),
        ("C", "B", "C − B"),
    ]
    labels: list[str] = []
    mean_diff: list[float] = []
    t_stats: list[float] = []
    p_unc: list[float] = []

    for hi, lo, label in pairs:
        d = wide[hi] - wide[lo]
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", category=RuntimeWarning)
            t_stat, p_u = stats.ttest_1samp(d, popmean=0.0, alternative="two-sided")
        labels.append(label)
        mean_diff.append(float(d.mean()))
        t_stats.append(float(t_stat) if np.isfinite(t_stat) else float("nan"))
        p_unc.append(float(p_u) if np.isfinite(p_u) else float("nan"))

    p_holm: list[float] = [float("nan")] * len(p_unc)
    finite_idx = [i for i, p in enumerate(p_unc) if np.isfinite(p)]
    if finite_idx:
        sub = [p_unc[i] for i in finite_idx]
        _, adj, _, _ = multipletests(sub, method="holm")
        for j, i in enumerate(finite_idx):
            p_holm[i] = float(adj[j])

    paras = [
        "Paired t-tests on the same scenarios across prompts, with Holm adjustment across the three contrasts.",
        "If the omnibus ANOVA is significant (p < 0.05), use p (Holm) < 0.05 to claim a pairwise difference.",
    ]
    headers = ["Contrast", "Mean Δ", "t", "p (uncorr.)", "p (Holm)"]
    rows: list[list[str]] = []
    for i, lab in enumerate(labels):
        md = mean_diff[i]
        ts = f"{t_stats[i]:.4f}" if np.isfinite(t_stats[i]) else "—"
        pu = f"{p_unc[i]:.4g}" if np.isfinite(p_unc[i]) else "—"
        ph = f"{p_holm[i]:.4g}" if np.isfinite(p_holm[i]) else "—"
        rows.append([lab, f"{md:.4f}", ts, pu, ph])
    return paras, headers, rows


def _regression_cluster_ols(
    long: pd.DataFrame,
) -> tuple[list[str], list[str], list[list[str]]]:
    """OLS Overall ~ Prompt (ref A); cluster-robust SE by ScenarioID."""
    d = long.copy()
    d["Prompt"] = pd.Categorical(d["Prompt"], categories=["A", "B", "C"], ordered=True)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        fit = smf.ols(
            'Overall ~ C(Prompt, Treatment(reference="A"))',
            data=d,
        ).fit(cov_type="cluster", cov_kwds={"groups": d["ScenarioID"]})

    idx = fit.params.index.tolist()
    term_labels: dict[str, str] = {}
    for name in idx:
        if name == "Intercept":
            term_labels[name] = "Intercept (Prompt A mean)"
        elif "[T.B]" in name:
            term_labels[name] = "Prompt B vs A"
        elif "[T.C]" in name:
            term_labels[name] = "Prompt C vs A"
        else:
            term_labels[name] = name

    ci = fit.conf_int()
    headers = ["Term", "Coef.", "SE", "z", "p", "95% CI low", "95% CI high"]
    rows: list[list[str]] = []
    for name in idx:
        coef = float(fit.params[name])
        se = float(fit.bse[name])
        pv = float(fit.pvalues[name])
        lo, hi = float(ci.loc[name, 0]), float(ci.loc[name, 1])
        z = coef / se if se > 0 else float("nan")
        rows.append(
            [
                term_labels.get(name, name),
                f"{coef:.4f}",
                f"{se:.4f}",
                f"{z:.3f}" if np.isfinite(z) else "—",
                f"{pv:.4g}" if pv >= 1e-6 else "<1e-6",
                f"{lo:.4f}",
                f"{hi:.4f}",
            ]
        )

    n_obs = int(fit.nobs)
    n_clust = int(d["ScenarioID"].nunique())
    r2 = float(fit.rsquared)
    paras = [
        "Linear regression of Overall score on prompt indicators (reference category: Prompt A). "
        "Standard errors are cluster-robust, clustering on scenario, so repeated observations from the "
        "same scenario are not treated as independent.",
        f"Model fit: R² = {r2:.3f}; N = {n_obs} observations in {n_clust} scenarios; covariance estimator: cluster-robust (HC0-type by cluster). "
        "With a balanced design, the intercept matches the mean for Prompt A; coefficients for B and C are the "
        "marginal mean differences versus A (consistent with the marginal means in the descriptives).",
    ]
    return paras, headers, rows


def _write_prompt_figure_files(long: pd.DataFrame, wide: pd.DataFrame, out_dir: Path) -> list[FigurePanel]:
    """Save boxplot, RM line plot, and mean±95% CI plot; return panels with paths."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    out_dir.mkdir(parents=True, exist_ok=True)
    order = ["A", "B", "C"]
    lo = long.copy()
    lo["Prompt"] = pd.Categorical(lo["Prompt"], categories=order, ordered=True)

    # --- 1. Boxplot ---
    fig, ax = plt.subplots(figsize=(6.2, 4.0), dpi=150)
    groups = [lo.loc[lo["Prompt"] == p, "Overall"].astype(float).to_numpy() for p in order]
    bp_kw: dict = {"patch_artist": True, "medianprops": {"color": "#023047", "linewidth": 1.5}}
    try:
        bp = ax.boxplot(groups, tick_labels=[f"Prompt {p}" for p in order], **bp_kw)
    except TypeError:
        bp = ax.boxplot(groups, labels=[f"Prompt {p}" for p in order], **bp_kw)
    colors = ("#8ecae6", "#219ebc", "#457b9d")
    for patch, color in zip(bp["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_alpha(0.75)
    ax.set_xlabel("Prompt condition")
    ax.set_ylabel("Overall score")
    ax.set_title("Overall score by prompt (boxplot)")
    ax.grid(True, axis="y", alpha=0.35, linestyle="--", linewidth=0.7)
    fig.tight_layout()
    p_box = out_dir / "fig_boxplot.png"
    fig.savefig(p_box, bbox_inches="tight")
    plt.close(fig)

    # --- 2. Repeated-measures line plot ---
    fig, ax = plt.subplots(figsize=(6.2, 4.0), dpi=150)
    xv = np.arange(3)
    cmap = plt.get_cmap("tab10")
    wide_s = wide.sort_index()
    for i, (_sid, row) in enumerate(wide_s.iterrows()):
        ys = [float(row[p]) for p in order]
        ax.plot(
            xv,
            ys,
            marker="o",
            markersize=4,
            linewidth=1.1,
            alpha=0.72,
            color=cmap(i % 10),
        )
    ax.set_xticks(xv, [f"Prompt {p}" for p in order])
    ax.set_xlabel("Prompt condition")
    ax.set_ylabel("Overall score")
    ax.set_title("Within-scenario profiles (one line per scenario)")
    ax.grid(True, alpha=0.35, linestyle="--", linewidth=0.7)
    fig.tight_layout()
    p_line = out_dir / "fig_lines_rm.png"
    fig.savefig(p_line, bbox_inches="tight")
    plt.close(fig)

    # --- 3. Mean ± 95% CI ---
    fig, ax = plt.subplots(figsize=(6.2, 4.0), dpi=150)
    means: list[float] = []
    ci_half: list[float] = []
    for p in order:
        v = lo.loc[lo["Prompt"] == p, "Overall"].astype(float).to_numpy()
        n = int(v.size)
        m = float(np.mean(v))
        sem = float(stats.sem(v)) if n > 1 else 0.0
        tcrit = float(stats.t.ppf(0.975, n - 1)) if n > 1 else 0.0
        means.append(m)
        ci_half.append(tcrit * sem)
    xv2 = np.arange(3)
    ax.errorbar(
        xv2,
        means,
        yerr=ci_half,
        fmt="o",
        markersize=9,
        capsize=7,
        color="#1d3557",
        ecolor="#457b9d",
        elinewidth=1.5,
        markerfacecolor="#e63946",
        markeredgecolor="#1d3557",
    )
    ax.set_xticks(xv2, [f"Prompt {p}" for p in order])
    ax.set_xlabel("Prompt condition")
    ax.set_ylabel("Mean overall score")
    ax.set_title("Marginal means ± 95% confidence intervals")
    ax.grid(True, axis="y", alpha=0.35, linestyle="--", linewidth=0.7)
    fig.tight_layout()
    p_ci = out_dir / "fig_mean_ci.png"
    fig.savefig(p_ci, bbox_inches="tight")
    plt.close(fig)

    box_paras = [
        "This is the primary distribution view for the experiment. It shows score distributions, medians, "
        "spread and variance, and overlap between Prompt A, B, and C. Use it alongside the repeated-measures ANOVA.",
        "Axes: horizontal axis = prompt condition (A, B, C); vertical axis = Overall score.",
        "What to look for: if the experimental pattern is A lowest, B intermediate, and C highest, you should "
        "see separation and ordering consistent with that pattern (medians and boxes shifted upward from A to C).",
    ]
    line_paras = [
        "Ideal for this design because the same scenarios are evaluated under each prompt. Each line is one scenario "
        "across Prompt A, B, and C.",
        "Axes: horizontal axis = prompt condition; vertical axis = Overall score.",
        "What to look for: many lines trending upward from A → B → C indicates within-scenario improvement as the "
        "prompt changes, which is strong evidence for a prompt effect at the paired level.",
    ]
    ci_paras = [
        "Shows marginal means with uncertainty (95% confidence intervals of the mean per prompt). This is a clean "
        "companion to the omnibus ANOVA and effect size.",
        "Axes: horizontal axis = prompt condition; vertical axis = mean Overall score. Error bars are 95% CIs of the mean "
        f"(n = {len(wide_s)} scenarios per prompt).",
        "Alternative: you could plot mean ± SD instead of CI for raw dispersion; CIs emphasize precision of the mean.",
    ]

    return [
        FigurePanel(
            "1. Boxplot (primary distribution comparison)",
            box_paras,
            p_box,
        ),
        FigurePanel(
            "2. Repeated-measures line plot",
            line_paras,
            p_line,
        ),
        FigurePanel(
            "3. Mean with 95% confidence intervals",
            ci_paras,
            p_ci,
        ),
    ]


def build_report_blocks(
    workbook_path: Path,
    wide: pd.DataFrame,
    long: pd.DataFrame,
    figure_panels: list[FigurePanel] | None = None,
) -> list[ReportBlock]:
    n = len(wide)
    intro_paras = [
        f"Workbook: {workbook_path.name}",
        f"Design: repeated measures — {n} scenarios × three prompts (A, B, C).",
        _strip_md_bold(
            "Outcome per row: numeric **Overall (column K)** when present; otherwise mean of D1–D6 (E–J). "
            "Re-save the workbook in Excel once if column K holds formulas and cached values look missing."
        ),
    ]

    desc_h, desc_r = _descriptive_table(long)
    anova_body, eta_body = _anova_effect_parts(long)
    post_paras, post_h, post_r = _posthoc_holm_parts(wide)
    reg_paras, reg_h, reg_r = _regression_cluster_ols(long)

    blocks: list[ReportBlock] = [
        ReportBlock("Summary", paragraphs=intro_paras),
        ReportBlock(
            "1. Descriptive statistics",
            paragraphs=["Mean and standard deviation of Overall by prompt (complete cases only)."],
            table_headers=desc_h,
            table_rows=desc_r,
        ),
    ]

    if figure_panels:
        blocks.append(
            ReportBlock(
                "2. Figures for repeated-measures design",
                paragraphs=[
                    "Three standard figures for a within-subject (repeated-measures) prompt comparison: "
                    "distributions, individual scenario trajectories, and marginal means with uncertainty. "
                    "Together they complement the ANOVA and post-hoc tests below.",
                ],
                panels=figure_panels,
                table_headers=["Plot", "Purpose"],
                table_rows=[
                    [
                        "Boxplot",
                        "Main distribution comparison across prompts (medians, spread, overlap); supports the RM ANOVA.",
                    ],
                    [
                        "Repeated-measures line plot",
                        "Each line is one scenario across A → B → C; upward trends show within-scenario improvement.",
                    ],
                    [
                        "Mean ± 95% CI",
                        "Marginal means with error bars (95% CI of the mean); concise summary of average differences.",
                    ],
                ],
            )
        )
        anova_title = "3. Repeated-measures ANOVA"
        eta_title = "4. Effect size (partial eta-squared)"
        post_title = "5. Post-hoc pairwise comparisons"
        reg_title = "6. Regression analysis"
    else:
        anova_title = "2. Repeated-measures ANOVA"
        eta_title = "3. Effect size (partial eta-squared)"
        post_title = "4. Post-hoc pairwise comparisons"
        reg_title = "5. Regression analysis"

    blocks.extend(
        [
            ReportBlock(
                anova_title,
                paragraphs=anova_body.replace("\r\n", "\n").strip().split("\n\n"),
            ),
            ReportBlock(
                eta_title,
                paragraphs=[p for p in eta_body.replace("\r\n", "\n").strip().split("\n") if p.strip()],
            ),
            ReportBlock(
                post_title,
                paragraphs=post_paras,
                table_headers=post_h,
                table_rows=post_r,
            ),
            ReportBlock(
                reg_title,
                paragraphs=reg_paras,
                table_headers=reg_h,
                table_rows=reg_r,
            ),
        ]
    )
    return blocks


def _set_cell_shading(cell, fill: str) -> None:
    """fill: hex e.g. 'D9E2F3' without #."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_body_paragraph(document: Document, text: str, *, italic: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = italic


def _add_data_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = document.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        hp = hdr[i].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(h)
        hr.bold = True
        hr.font.name = "Calibri"
        hr.font.size = Pt(10)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(hdr[i], "D9E2F3")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def write_docx(out_path: Path, workbook_path: Path, blocks: list[ReportBlock]) -> None:
    document = Document()
    title = document.add_heading("AI Report Validation System — statistical report", level=0)
    for r in title.runs:
        r.font.name = "Calibri Light"
        r.font.size = Pt(26)
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r1 = sub.add_run(f"Source workbook: {workbook_path.name}")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.add_run("\n")
    r2 = sub.add_run(
        "AI Report Validation System — descriptives, figures (boxplot, RM line plot, mean±95% CI), "
        "repeated-measures ANOVA, Holm-adjusted pairwise tests, and cluster-robust regression."
    )
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for i, block in enumerate(blocks):
        h = document.add_heading(block.title, level=1)
        h.paragraph_format.space_before = Pt(10 if i == 0 else 18)
        h.paragraph_format.space_after = Pt(6)
        for r in h.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        for para in block.paragraphs:
            if para.strip():
                _add_body_paragraph(document, para.strip())
        for panel in block.panels:
            h2 = document.add_heading(panel.heading, level=2)
            h2.paragraph_format.space_before = Pt(12)
            h2.paragraph_format.space_after = Pt(4)
            for r in h2.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x1F, 0x5C, 0x7A)
            for para in panel.paragraphs:
                if para.strip():
                    _add_body_paragraph(document, para.strip())
            if panel.image_path.is_file():
                document.add_picture(str(panel.image_path), width=Inches(6.2))
            document.add_paragraph()
        if block.table_headers and block.table_rows:
            _add_data_table(document, block.table_headers, block.table_rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def run_report(
    workbook: Path,
    docx_path: Path | None,
    stdout_only: bool,
    csv_out: Path | None,
) -> int:
    if not workbook.is_file():
        print(f"Workbook not found: {workbook}", file=sys.stderr)
        return 1

    try:
        df = build_long_frame(workbook)
    except ValueError as e:
        print(e, file=sys.stderr)
        return 1

    if csv_out:
        df.to_csv(csv_out, index=False)

    try:
        wide, long = _wide_complete(df)
    except ValueError as e:
        print(e, file=sys.stderr)
        return 1

    if stdout_only:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            blocks = build_report_blocks(workbook, wide, long, figure_panels=None)
        buf = io.StringIO()
        for block in blocks:
            buf.write(f"=== {block.title} ===\n")
            for para in block.paragraphs:
                buf.write(f"{para}\n")
            for panel in block.panels:
                buf.write(f"\n--- {panel.heading} ---\n")
                for para in panel.paragraphs:
                    buf.write(f"{para}\n")
                buf.write("(Figure omitted in stdout mode; run without --stdout-only to embed in .docx.)\n")
            if block.table_headers and block.table_rows:
                buf.write("\n")
                buf.write(_ascii_table(block.table_headers, block.table_rows))
                buf.write("\n")
            buf.write("\n")
        sys.stdout.write(buf.getvalue())
        if csv_out:
            sys.stdout.write(f"(CSV written to {csv_out})\n")
        return 0

    assert docx_path is not None
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            panels = _write_prompt_figure_files(long, wide, tmp_path)
            blocks = build_report_blocks(workbook, wide, long, figure_panels=panels)
            write_docx(docx_path, workbook, blocks)
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--workbook", type=Path, default=DEFAULT_XLSX)
    ap.add_argument(
        "--docx",
        type=Path,
        default=None,
        help="Output .docx (default: <workbook_stem>_significance_report.docx beside workbook)",
    )
    ap.add_argument(
        "--stdout-only",
        action="store_true",
        help="Print report to stdout instead of Word",
    )
    ap.add_argument("--csv-out", type=Path, default=None)
    args = ap.parse_args()

    wb = args.workbook
    if args.stdout_only:
        return run_report(wb, None, True, args.csv_out)

    docx_path = args.docx or wb.with_name(f"{wb.stem}_significance_report.docx")
    return run_report(wb, docx_path, False, args.csv_out)


if __name__ == "__main__":
    raise SystemExit(main())
